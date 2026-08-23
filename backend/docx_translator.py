import docx
from pptx_translator import translate_text_batch
import concurrent.futures

def extract_and_translate_docx(input_path: str, output_path: str, target_language: str, model_name: str = "openai/gpt-5.6-luna"):
    """
    Translates a Microsoft Word document into a target language while preserving run-level formatting
    by capturing the format of the first run of a paragraph and applying it to the translated text.
    """
    doc = docx.Document(input_path)
    
    # Collect all paragraphs that have text (from body and tables)
    paragraphs_to_translate = []
    
    # Body paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            paragraphs_to_translate.append(paragraph)
            
    # Table paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        paragraphs_to_translate.append(paragraph)
                        
    # Extract strings
    original_texts = [p.text for p in paragraphs_to_translate]
    
    # Smart batching by character count to minimize API calls
    # We use a very small limit (800 chars) so the LLM responds extremely fast,
    # preventing the API proxy from timing out and throwing a 500 error.
    batches = []
    current_batch = []
    current_length = 0
    
    for text in original_texts:
        if current_length + len(text) > 800 and current_batch:
            batches.append(current_batch)
            current_batch = []
            current_length = 0
            
        current_batch.append(text)
        current_length += len(text)
        
    if current_batch:
        batches.append(current_batch)
        
    translated_texts = []
    
    # Process sequentially to avoid concurrent proxy crashes
    for batch in batches:
        translated_texts.extend(translate_text_batch(batch, target_language, model_name))
            
    # Safely apply translations back
    for i, paragraph in enumerate(paragraphs_to_translate):
        if i < len(translated_texts):
            translated_text = translated_texts[i]
            
            # Save the formatting of the first run (if any exist)
            font_name = font_size = font_bold = font_italic = font_underline = font_color = None
            if paragraph.runs:
                first_run = paragraph.runs[0]
                try: font_name = first_run.font.name
                except: pass
                
                try: font_size = first_run.font.size
                except: pass
                
                try: font_bold = first_run.font.bold
                except: pass
                
                try: font_italic = first_run.font.italic
                except: pass
                
                try: font_underline = first_run.font.underline
                except: pass
                
                try: font_color = first_run.font.color.rgb if hasattr(first_run.font, 'color') and hasattr(first_run.font.color, 'rgb') else None
                except: pass

            # Clear existing runs
            paragraph.clear()
            
            # Add a new single run with the translated text
            new_run = paragraph.add_run(translated_text)
            
            # Apply saved formatting (intentionally skipping font_name to allow Word's font fallback for non-Latin scripts)
            if font_size is not None: new_run.font.size = font_size
            if font_bold is not None: new_run.font.bold = font_bold
            if font_italic is not None: new_run.font.italic = font_italic
            if font_underline is not None: new_run.font.underline = font_underline
            if font_color is not None: 
                try:
                    new_run.font.color.rgb = font_color
                except:
                    pass

    # Save the modified document
    doc.save(output_path)
    return output_path
